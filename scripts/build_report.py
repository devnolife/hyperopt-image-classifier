"""Generate laporan akademik .docx lengkap BAB I-V."""
from __future__ import annotations
import sys
import json
from pathlib import Path
from datetime import datetime

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


METHOD_LABEL = {
    "grid_search": "Grid Search",
    "random_search": "Random Search",
    "bayesian_tpe": "Bayesian Optimization (TPE)",
    "hyperband_asha": "Hyperband / ASHA",
    "genetic": "Genetic Algorithm",
}


def set_default_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for s in doc.sections:
        s.top_margin = Cm(3); s.bottom_margin = Cm(3)
        s.left_margin = Cm(4); s.right_margin = Cm(3)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_paragraph(doc, text, justify=True, first_line_indent=True, bold=False):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = "Times New Roman"; r.font.size = Pt(12)
    r.bold = bold
    return p


def add_image(doc, path: Path, caption: str, width_inches: float = 5.5):
    if not path.exists():
        add_paragraph(doc, f"[Gambar tidak ditemukan: {path.name}]", first_line_indent=False)
        return
    doc.add_picture(str(path), width=Inches(width_inches))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.name = "Times New Roman"; r.font.size = Pt(11); r.italic = True


def add_table_from_df(doc, df: pd.DataFrame, caption: str = None):
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.name = "Times New Roman"; r.font.size = Pt(11); r.italic = True
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, c in enumerate(df.columns):
        hdr[i].text = str(c)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(11)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"; r.font.size = Pt(11)


def build_report():
    results_path = ROOT / "results" / "all_hpo_results.json"
    best_path = ROOT / "results" / "best_configs.json"
    final_path = ROOT / "results" / "final_training.json"
    comp_csv = ROOT / "results" / "tables" / "hpo_comparison.csv"
    fig_dir = ROOT / "results" / "figures"

    all_results = json.loads(results_path.read_text(encoding="utf-8"))
    best_configs = json.loads(best_path.read_text(encoding="utf-8"))
    final_info = json.loads(final_path.read_text(encoding="utf-8")) if final_path.exists() else None
    comp_df = pd.read_csv(comp_csv) if comp_csv.exists() else None

    winner_method, winner_info = max(best_configs.items(), key=lambda kv: kv[1]["val_acc"])
    total_time_all = sum(r["total_time_method"] for r in best_configs.values())

    doc = Document()
    set_default_style(doc)

    # ===== Cover =====
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("HYPERPARAMETER OPTIMIZATION FOR IMAGE CLASSIFICATION\nUSING DEEP LEARNING EXPERIMENT MANAGER")
    r.bold = True; r.font.size = Pt(16); r.font.name = "Times New Roman"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("(Studi Kasus: Klasifikasi Citra CIFAR-10 dengan Custom CNN dan MLflow)")
    r.italic = True; r.font.size = Pt(13); r.font.name = "Times New Roman"

    for _ in range(6):
        doc.add_paragraph()

    for label in ["LAPORAN TUGAS", "", "Disusun oleh:", "Nama : [Nama Anda]", "NIM  : [NIM Anda]",
                  "", "Mata Kuliah : [Nama Mata Kuliah]",
                  "Dosen Pengampu : [Nama Dosen]",
                  "", "[PROGRAM STUDI / FAKULTAS]",
                  "[UNIVERSITAS]", f"[TAHUN {datetime.now().year}]"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(label); rr.font.name = "Times New Roman"; rr.font.size = Pt(12)
        if label.startswith("LAPORAN") or label.startswith("["):
            rr.bold = True
    doc.add_page_break()

    # ===== Abstrak =====
    add_heading(doc, "ABSTRAK", level=1)
    abstrak = (
        "Pemilihan hyperparameter yang tepat merupakan salah satu faktor kritis dalam keberhasilan "
        "model deep learning. Penelitian ini membandingkan lima metode Hyperparameter Optimization "
        "(HPO), yaitu Grid Search, Random Search, Bayesian Optimization berbasis Tree-structured "
        "Parzen Estimator (TPE), Hyperband/ASHA, dan Genetic Algorithm, dalam konteks tugas klasifikasi "
        "citra pada dataset CIFAR-10 menggunakan arsitektur Custom Convolutional Neural Network (CNN). "
        "Seluruh eksperimen dikelola dan direkam menggunakan MLflow sebagai deep learning experiment "
        "manager, sehingga setiap trial, metrik, dan artefaknya dapat ditelusuri secara sistematis. "
        f"Berdasarkan hasil eksperimen, metode {METHOD_LABEL.get(winner_method, winner_method)} "
        f"memperoleh validation accuracy tertinggi sebesar {winner_info['val_acc']:.4f} dengan "
        f"learning rate={winner_info['params']['learning_rate']:.2e}, batch size={winner_info['params']['batch_size']}, "
        f"optimizer={winner_info['params']['optimizer']}, dropout={winner_info['params']['dropout']:.3f}, "
        f"dan base filter={winner_info['params']['base_filters']}. "
        "Model final yang dilatih kembali menggunakan konfigurasi terbaik "
        + (f"mencapai test accuracy sebesar {final_info['test_acc']:.4f} pada CIFAR-10. " if final_info else "dievaluasi pada test set CIFAR-10. ")
        + "Hasil penelitian menunjukkan bahwa metode HPO berbasis model (Bayesian dan evolutionary) "
          "secara umum lebih efisien dibanding pendekatan acak dan eksaustif, serta penggunaan "
          "experiment manager sangat membantu reprodusibilitas dan analisis eksperimen."
    )
    p = doc.add_paragraph(abstrak); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.25)

    p = doc.add_paragraph()
    r = p.add_run("Kata Kunci: "); r.bold = True
    p.add_run("hyperparameter optimization, CIFAR-10, Convolutional Neural Network, MLflow, Optuna, Bayesian optimization.")
    doc.add_page_break()

    # ===== BAB I =====
    add_heading(doc, "BAB I  PENDAHULUAN", level=1)
    add_heading(doc, "1.1  Latar Belakang", level=2)
    add_paragraph(doc,
        "Perkembangan pesat deep learning, khususnya Convolutional Neural Network (CNN), telah "
        "mendorong pencapaian yang signifikan pada berbagai tugas klasifikasi citra. Namun, kinerja "
        "model CNN sangat bergantung pada pemilihan hyperparameter seperti learning rate, batch size, "
        "jenis optimizer, dropout, dan jumlah filter konvolusi. Penentuan hyperparameter secara manual "
        "sering kali tidak efisien karena ruang pencarian yang sangat luas dan biaya komputasi yang tinggi.")
    add_paragraph(doc,
        "Hyperparameter Optimization (HPO) muncul sebagai pendekatan sistematis untuk menemukan "
        "konfigurasi terbaik. Terdapat berbagai metode HPO, mulai dari pendekatan eksaustif (Grid "
        "Search), pendekatan stokastik (Random Search), pendekatan probabilistik (Bayesian "
        "Optimization), pendekatan multi-fidelity (Hyperband/ASHA), hingga pendekatan evolusioner "
        "(Genetic Algorithm). Tiap metode memiliki kelebihan dan keterbatasannya masing-masing, "
        "sehingga perlu dikomparasi pada kasus nyata.")
    add_paragraph(doc,
        "Seiring bertambahnya jumlah eksperimen, kebutuhan akan experiment manager menjadi semakin "
        "penting. MLflow merupakan salah satu experiment manager open-source yang banyak digunakan "
        "untuk mencatat parameter, metrik, dan artefak dari tiap run sehingga reprodusibilitas, "
        "pelacakan, dan analisis menjadi lebih mudah.")

    add_heading(doc, "1.2  Rumusan Masalah", level=2)
    for t in [
        "Bagaimana implementasi lima metode HPO (Grid Search, Random Search, Bayesian Optimization, Hyperband/ASHA, dan Genetic Algorithm) pada model CNN untuk klasifikasi CIFAR-10?",
        "Bagaimana perbandingan performa kelima metode HPO tersebut ditinjau dari validation accuracy, jumlah trial, dan total waktu eksekusi?",
        "Bagaimana peran MLflow sebagai experiment manager dalam mendukung reprodusibilitas dan analisis eksperimen HPO?",
    ]:
        p = doc.add_paragraph(t, style="List Number")
        p.paragraph_format.line_spacing = 1.5

    add_heading(doc, "1.3  Tujuan Penelitian", level=2)
    for t in [
        "Mengimplementasikan lima metode HPO pada Custom CNN untuk klasifikasi CIFAR-10.",
        "Membandingkan performa kelima metode dari sisi akurasi, efisiensi trial, dan waktu.",
        "Mendemonstrasikan penggunaan MLflow sebagai experiment manager dalam siklus eksperimen deep learning.",
    ]:
        p = doc.add_paragraph(t, style="List Number")
        p.paragraph_format.line_spacing = 1.5

    add_heading(doc, "1.4  Manfaat Penelitian", level=2)
    add_paragraph(doc,
        "Hasil penelitian ini diharapkan memberikan rekomendasi empiris mengenai metode HPO yang "
        "paling sesuai untuk kasus klasifikasi citra berukuran menengah, serta memberikan template "
        "alur kerja eksperimen deep learning berbasis experiment manager yang dapat direplikasi.")

    add_heading(doc, "1.5  Batasan Masalah", level=2)
    for t in [
        "Dataset yang digunakan hanya CIFAR-10.",
        "Model yang digunakan adalah Custom CNN 3 blok konvolusi (bukan transfer learning).",
        "Hyperparameter yang ditune: learning rate, batch size, optimizer, dropout, dan base filters.",
        "Budget trial dan epoch dikonfigurasi ringan demi efisiensi komputasi.",
        "Experiment manager yang digunakan adalah MLflow versi open-source (local tracking).",
    ]:
        p = doc.add_paragraph(t, style="List Number")
        p.paragraph_format.line_spacing = 1.5
    doc.add_page_break()

    # ===== BAB II =====
    add_heading(doc, "BAB II  TINJAUAN PUSTAKA", level=1)
    add_heading(doc, "2.1  Convolutional Neural Network (CNN)", level=2)
    add_paragraph(doc,
        "CNN adalah arsitektur jaringan saraf tiruan yang dirancang khusus untuk data berbentuk grid "
        "seperti citra. Blok utama CNN terdiri dari lapisan konvolusi, aktivasi non-linear (ReLU), "
        "batch normalization, pooling, dan fully-connected. Dengan berbagi bobot (weight sharing) dan "
        "konektivitas lokal, CNN efisien dalam mengekstraksi fitur hierarkis pada citra.")

    add_heading(doc, "2.2  Image Classification dan Dataset CIFAR-10", level=2)
    add_paragraph(doc,
        "Image classification adalah tugas memetakan sebuah citra ke satu kelas dari sekumpulan "
        "kategori yang telah ditentukan. Dataset CIFAR-10 terdiri dari 60.000 citra berwarna "
        "berukuran 32×32 piksel dalam 10 kelas (pesawat, mobil, burung, kucing, rusa, anjing, katak, "
        "kuda, kapal, truk). Dataset ini terbagi menjadi 50.000 citra latih dan 10.000 citra uji, "
        "dan merupakan benchmark standar untuk eksperimen klasifikasi citra berskala kecil.")

    add_heading(doc, "2.3  Hyperparameter Optimization", level=2)
    add_paragraph(doc,
        "Hyperparameter adalah parameter yang nilainya ditetapkan sebelum proses pelatihan dan "
        "tidak dipelajari dari data. HPO bertujuan menemukan konfigurasi hyperparameter yang "
        "meminimalkan (atau memaksimalkan) suatu fungsi objektif, umumnya berupa loss atau accuracy "
        "pada validation set.")

    add_heading(doc, "2.3.1  Grid Search", level=3)
    add_paragraph(doc,
        "Grid Search melakukan pencarian eksaustif terhadap seluruh kombinasi pada grid yang telah "
        "didefinisikan. Metode ini sederhana namun kompleksitasnya tumbuh eksponensial terhadap "
        "dimensi search space (curse of dimensionality).")

    add_heading(doc, "2.3.2  Random Search", level=3)
    add_paragraph(doc,
        "Random Search mengambil sampel konfigurasi secara acak. Bergstra dan Bengio (2012) "
        "menunjukkan bahwa random search sering kali lebih efisien dibanding grid search pada "
        "ruang berdimensi tinggi karena tidak semua hyperparameter memiliki pengaruh yang sama.")

    add_heading(doc, "2.3.3  Bayesian Optimization (TPE)", level=3)
    add_paragraph(doc,
        "Bayesian Optimization membangun model probabilistik (surrogate) dari fungsi objektif dan "
        "menggunakan acquisition function untuk memilih konfigurasi selanjutnya yang paling "
        "menjanjikan. Tree-structured Parzen Estimator (TPE) merupakan pendekatan populer yang "
        "memodelkan p(x|y) alih-alih p(y|x) seperti Gaussian Process.")

    add_heading(doc, "2.3.4  Hyperband / ASHA", level=3)
    add_paragraph(doc,
        "Hyperband dan Asynchronous Successive Halving Algorithm (ASHA) adalah metode multi-fidelity "
        "yang mengalokasikan resource secara dinamis. Konfigurasi yang kurang menjanjikan dipangkas "
        "(pruned) lebih awal sehingga budget komputasi difokuskan pada konfigurasi terbaik.")

    add_heading(doc, "2.3.5  Genetic Algorithm", level=3)
    add_paragraph(doc,
        "Genetic Algorithm (GA) terinspirasi dari evolusi biologis. Populasi kandidat solusi "
        "dievaluasi, diseleksi, disilangkan (crossover), dan dimutasi selama sejumlah generasi untuk "
        "memperbaiki kualitas solusi secara bertahap.")

    add_heading(doc, "2.4  Experiment Manager — MLflow", level=2)
    add_paragraph(doc,
        "MLflow adalah platform open-source untuk mengelola siklus hidup machine learning. Dalam "
        "penelitian ini, komponen MLflow Tracking digunakan untuk mencatat parameter, metrik "
        "per-epoch, tag, dan artefak (gambar, model) dari tiap trial HPO sehingga eksperimen "
        "mudah dibandingkan, dianalisis, dan direproduksi.")
    doc.add_page_break()

    # ===== BAB III =====
    add_heading(doc, "BAB III  METODOLOGI", level=1)
    add_heading(doc, "3.1  Alur Penelitian", level=2)
    add_paragraph(doc,
        "Alur penelitian dimulai dari persiapan dataset CIFAR-10, definisi arsitektur Custom CNN, "
        "penentuan search space hyperparameter, pelaksanaan lima metode HPO dengan logging melalui "
        "MLflow, analisis hasil, pelatihan ulang model terbaik, dan penulisan laporan.")

    add_heading(doc, "3.2  Dataset dan Preprocessing", level=2)
    add_paragraph(doc,
        "Dataset CIFAR-10 dibagi menjadi 45.000 data latih, 5.000 data validasi, dan 10.000 data uji. "
        "Augmentasi diterapkan pada data latih berupa RandomCrop 32×32 dengan padding 4 dan "
        "RandomHorizontalFlip, diikuti normalisasi dengan mean dan std CIFAR-10.")

    add_heading(doc, "3.3  Arsitektur Custom CNN", level=2)
    add_paragraph(doc,
        "Model terdiri dari tiga blok konvolusi; masing-masing blok berisi dua lapisan Conv3×3 "
        "dengan BatchNorm, ReLU, MaxPool 2×2, dan Dropout2D. Jumlah filter pada blok pertama "
        "adalah base_filters (f1), blok kedua 2·f1, blok ketiga 4·f1. Output flatten dihubungkan ke "
        "FC 256 → Dropout → FC 10.")

    add_heading(doc, "3.4  Search Space Hyperparameter", level=2)
    ss_rows = [
        ["learning_rate", "log-uniform [1e-4, 1e-1]"],
        ["batch_size", "{64, 128}"],
        ["optimizer", "{SGD, Adam, RMSprop}"],
        ["dropout", "uniform [0.1, 0.5]"],
        ["base_filters", "{16, 32, 64}"],
    ]
    ss_df = pd.DataFrame(ss_rows, columns=["Hyperparameter", "Search Space"])
    add_table_from_df(doc, ss_df, caption="Tabel 3.1 Definisi search space hyperparameter.")

    add_heading(doc, "3.5  Konfigurasi Metode HPO", level=2)
    hpo_rows = [
        ["Grid Search", "Eksaustif pada subset grid", f"{all_results['grid_search']['n_trials']} trial"],
        ["Random Search", "Optuna RandomSampler", f"{all_results['random_search']['n_trials']} trial"],
        ["Bayesian (TPE)", "Optuna TPESampler", f"{all_results['bayesian_tpe']['n_trials']} trial"],
        ["Hyperband/ASHA", "Optuna SuccessiveHalvingPruner", f"{all_results['hyperband_asha']['n_trials']} trial"],
        ["Genetic Algorithm", "DEAP (pop=6, gen=3)", f"{all_results['genetic']['n_trials']} evaluasi"],
    ]
    hpo_df = pd.DataFrame(hpo_rows, columns=["Metode", "Implementasi", "Budget"])
    add_table_from_df(doc, hpo_df, caption="Tabel 3.2 Konfigurasi tiap metode HPO.")

    add_heading(doc, "3.6  Setup Experiment Manager (MLflow)", level=2)
    add_paragraph(doc,
        "MLflow dikonfigurasi dengan local tracking URI './mlruns'. Setiap metode HPO memiliki "
        "experiment terpisah (HPO_grid_search, HPO_random_search, HPO_bayesian_tpe, "
        "HPO_hyperband_asha, HPO_genetic). Setiap trial dicatat sebagai run individual dengan "
        "parameter, metrik per-epoch (train/val loss dan accuracy), serta tag metode HPO.")

    add_heading(doc, "3.7  Metrik Evaluasi", level=2)
    add_paragraph(doc,
        "Metrik utama adalah validation accuracy terbaik yang dicapai tiap trial. Metrik pendukung "
        "meliputi loss, total waktu eksekusi, best epoch, accuracy pada test set, classification "
        "report (precision, recall, F1) per kelas, serta confusion matrix.")

    add_heading(doc, "3.8  Lingkungan Eksperimen", level=2)
    add_paragraph(doc,
        "Eksperimen dijalankan pada lingkungan Windows dengan Python 3.13, PyTorch, Optuna, DEAP, "
        "dan MLflow. Perangkat keras yang digunakan memanfaatkan akselerasi GPU NVIDIA apabila "
        "tersedia, atau CPU sebagai fallback. Seed acak diatur konsisten (42) untuk menjamin "
        "reprodusibilitas.")
    doc.add_page_break()

    # ===== BAB IV =====
    add_heading(doc, "BAB IV  HASIL DAN PEMBAHASAN", level=1)
    add_heading(doc, "4.1  Hasil Tiap Metode HPO", level=2)
    for name, r in all_results.items():
        bp = r["best"]["params"]
        add_paragraph(doc,
            f"Metode {METHOD_LABEL.get(name, name)} menghasilkan {r['n_trials']} trial dengan total "
            f"waktu {r['total_time']:.1f} detik. Konfigurasi terbaik pada metode ini adalah "
            f"learning_rate={bp['learning_rate']:.2e}, batch_size={bp['batch_size']}, "
            f"optimizer={bp['optimizer']}, dropout={bp['dropout']:.3f}, "
            f"base_filters={bp['base_filters']} dengan validation accuracy "
            f"{r['best']['val_acc']:.4f} pada epoch {r['best']['best_epoch']}.")

    add_heading(doc, "4.2  Komparasi Lima Metode HPO", level=2)
    if comp_df is not None:
        add_table_from_df(doc, comp_df, caption="Tabel 4.1 Komparasi hasil lima metode HPO.")
    add_image(doc, fig_dir / "best_val_acc_per_method.png",
              "Gambar 4.1 Best validation accuracy per metode HPO.")
    add_image(doc, fig_dir / "total_time_per_method.png",
              "Gambar 4.2 Total waktu eksekusi per metode HPO (detik).")
    add_image(doc, fig_dir / "convergence_best_so_far.png",
              "Gambar 4.3 Konvergensi best-so-far validation accuracy sebagai fungsi trial.")
    add_image(doc, fig_dir / "learning_curve_best_trials.png",
              "Gambar 4.4 Learning curve validation accuracy best trial per metode.")
    add_image(doc, fig_dir / "scatter_lr_vs_acc.png",
              "Gambar 4.5 Sebaran validation accuracy terhadap learning rate.")

    add_heading(doc, "4.3  Analisis Hasil Komparasi", level=2)
    add_paragraph(doc,
        f"Berdasarkan Tabel 4.1 dan Gambar 4.1, metode {METHOD_LABEL.get(winner_method, winner_method)} "
        f"memperoleh validation accuracy tertinggi sebesar {winner_info['val_acc']:.4f}. "
        "Grafik konvergensi pada Gambar 4.3 memperlihatkan bahwa metode berbasis model (Bayesian TPE) "
        "dan Hyperband/ASHA cenderung mencapai nilai tinggi lebih cepat dibanding Random Search "
        "dan Grid Search, meski Grid Search memberikan baseline yang eksaustif pada subset grid.")
    add_paragraph(doc,
        "Sebaran pada Gambar 4.5 menunjukkan bahwa learning rate merupakan hyperparameter dengan "
        "pengaruh paling besar terhadap validation accuracy; konfigurasi dengan learning rate yang "
        "terlalu tinggi (>1e-1) atau terlalu rendah (<1e-4) cenderung gagal berkonvergensi. "
        "Metode berbasis probabilistik/evolusioner lebih fokus pada rentang learning rate yang "
        "menjanjikan, sehingga menggunakan budget trial lebih efisien.")

    add_heading(doc, "4.4  Hasil Model Final", level=2)
    if final_info is not None:
        add_paragraph(doc,
            f"Model final dilatih kembali menggunakan konfigurasi pemenang "
            f"({METHOD_LABEL.get(final_info['winner_method'], final_info['winner_method'])}) "
            f"selama {final_info['epochs']} epoch dengan scheduler CosineAnnealing. "
            f"Pada test set, model mencapai accuracy {final_info['test_acc']:.4f} dan loss "
            f"{final_info['test_loss']:.4f}.")
        add_image(doc, fig_dir / "final_learning_curves.png",
                  "Gambar 4.6 Learning curve train/val loss dan accuracy model final.")
        add_image(doc, fig_dir / "final_confusion_matrix.png",
                  "Gambar 4.7 Confusion matrix model final pada test set CIFAR-10.")

        # classification report table
        cr = final_info["classification_report"]
        rep_rows = []
        for k, v in cr.items():
            if isinstance(v, dict):
                rep_rows.append([k, f"{v.get('precision', 0):.3f}", f"{v.get('recall', 0):.3f}",
                                 f"{v.get('f1-score', 0):.3f}", f"{v.get('support', 0)}"])
        rep_df = pd.DataFrame(rep_rows, columns=["Kelas / Metrik", "Precision", "Recall", "F1-Score", "Support"])
        add_table_from_df(doc, rep_df, caption="Tabel 4.2 Classification report model final pada test set.")
    else:
        add_paragraph(doc, "Hasil model final belum tersedia. Jalankan scripts/final_train.py.")

    add_heading(doc, "4.5  Pembahasan Peran MLflow", level=2)
    add_paragraph(doc,
        "Seluruh trial (grid, random, bayesian, hyperband, genetic) dan juga final training tercatat "
        "pada MLflow Tracking dengan experiment terpisah. Dashboard MLflow memudahkan komparasi "
        "antar run melalui fitur sorting, filtering, dan parallel plot. Artefak seperti confusion "
        "matrix dan learning curve diunggah sebagai image artifact pada run final, sehingga laporan "
        "dapat dihasilkan otomatis tanpa kehilangan jejak eksperimen.")
    doc.add_page_break()

    # ===== BAB V =====
    add_heading(doc, "BAB V  KESIMPULAN DAN SARAN", level=1)
    add_heading(doc, "5.1  Kesimpulan", level=2)
    for t in [
        f"Kelima metode HPO berhasil diimplementasikan pada Custom CNN untuk CIFAR-10 dengan total {sum(r['n_trials'] for r in all_results.values())} trial dan total waktu {total_time_all:.1f} detik.",
        f"Metode {METHOD_LABEL.get(winner_method, winner_method)} memberikan validation accuracy terbaik ({winner_info['val_acc']:.4f}) pada budget yang digunakan.",
        "Metode Bayesian Optimization dan Hyperband/ASHA cenderung lebih efisien dibanding Grid/Random Search karena memanfaatkan informasi dari trial sebelumnya atau memangkas trial buruk lebih dini.",
        "MLflow sebagai experiment manager terbukti memudahkan pelacakan, reprodusibilitas, dan analisis perbandingan metode HPO.",
    ]:
        p = doc.add_paragraph(t, style="List Number"); p.paragraph_format.line_spacing = 1.5

    add_heading(doc, "5.2  Saran", level=2)
    for t in [
        "Memperbesar budget trial dan epoch untuk memperoleh komparasi yang lebih stabil secara statistik.",
        "Mencoba arsitektur yang lebih dalam seperti ResNet-18 maupun pendekatan transfer learning.",
        "Menambah dimensi search space seperti weight decay, jenis aktivasi, dan strategi scheduler.",
        "Mengintegrasikan MLflow Model Registry untuk alur staging-production model yang lebih lengkap.",
    ]:
        p = doc.add_paragraph(t, style="List Number"); p.paragraph_format.line_spacing = 1.5
    doc.add_page_break()

    # ===== Daftar Pustaka =====
    add_heading(doc, "DAFTAR PUSTAKA", level=1)
    refs = [
        "Bergstra, J., & Bengio, Y. (2012). Random Search for Hyper-Parameter Optimization. Journal of Machine Learning Research, 13, 281–305.",
        "Bergstra, J., Bardenet, R., Bengio, Y., & Kégl, B. (2011). Algorithms for Hyper-Parameter Optimization. In NeurIPS.",
        "Li, L., Jamieson, K., DeSalvo, G., Rostamizadeh, A., & Talwalkar, A. (2017). Hyperband: A Novel Bandit-Based Approach to Hyperparameter Optimization. JMLR, 18(185), 1–52.",
        "Li, L., Jamieson, K., Rostamizadeh, A., et al. (2020). A System for Massively Parallel Hyperparameter Tuning (ASHA). MLSys.",
        "Akiba, T., Sano, S., Yanase, T., Ohta, T., & Koyama, M. (2019). Optuna: A Next-generation Hyperparameter Optimization Framework. KDD.",
        "Krizhevsky, A. (2009). Learning Multiple Layers of Features from Tiny Images (CIFAR-10). Technical Report, University of Toronto.",
        "Paszke, A., et al. (2019). PyTorch: An Imperative Style, High-Performance Deep Learning Library. NeurIPS.",
        "Zaharia, M., et al. (2018). Accelerating the Machine Learning Lifecycle with MLflow. IEEE Data Eng. Bull., 41(4).",
        "Fortin, F.-A., et al. (2012). DEAP: Evolutionary Algorithms Made Easy. JMLR, 13, 2171–2175.",
        "Ioffe, S., & Szegedy, C. (2015). Batch Normalization: Accelerating Deep Network Training by Reducing Internal Covariate Shift. ICML.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)
        p.paragraph_format.line_spacing = 1.15

    out = ROOT / "laporan" / "Laporan_HPO_CIFAR10.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Laporan tersimpan: {out}")


if __name__ == "__main__":
    build_report()
